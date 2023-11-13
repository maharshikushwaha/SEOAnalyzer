import requests
from bs4 import BeautifulSoup
from docx import Document

GOOGLE_MOBILE_FRIENDLY_API = "https://searchconsole.googleapis.com/v1/urlTestingTools/mobileFriendlyTest:run"

def get_html_content(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.text
    except requests.RequestException as e:
        return None

def get_page_title_and_meta_description(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    title_tag = soup.find('title')
    page_title = title_tag.get_text() if title_tag else 'Title not found'
    
    meta_description_tag = soup.find('meta', attrs={'name': 'description'})
    meta_description = meta_description_tag['content'] if meta_description_tag else 'Meta description not found'
    
    return page_title, meta_description

def check_keyword_presence(html_content, keywords):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    keyword_presence = {keyword: False for keyword in keywords}
    
    for keyword in keywords:
        if soup.body and soup.body.find(text=lambda text: keyword.lower() in text.lower()):
            keyword_presence[keyword] = True
    
    return keyword_presence

def measure_page_load_speed(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.elapsed.total_seconds()
    except requests.RequestException as e:
        return None

def check_mobile_friendly(url):
    try:
        payload = {'url': url}
        response = requests.post(GOOGLE_MOBILE_FRIENDLY_API, json=payload)
        response.raise_for_status()
        result = response.json()

        if 'mobileFriendliness' in result:
            return result['mobileFriendliness'] == 'MOBILE_FRIENDLY'
        else:
            return None
    except requests.RequestException as e:
        return None

def generate_report(url, keywords, page_title, meta_description, keyword_presence, page_load_speed, mobile_friendly):
    document = Document()
    document.add_heading('SEO Report', level=1)
    
    document.add_heading('1. Page Title', level=2)
    document.add_paragraph(page_title)
    
    document.add_heading('2. Meta Description', level=2)
    document.add_paragraph(meta_description)
    
    document.add_heading('3. Keywords Presence', level=2)
    for keyword, found in keyword_presence.items():
        document.add_paragraph(f"{keyword.strip()}: {'Found' if found else 'Not Found'}")
    
    document.add_heading('4. Page Load Speed', level=2)
    document.add_paragraph(f"{page_load_speed:.2f} seconds" if page_load_speed else "Not Available")
    
    document.add_heading('5. Mobile-Friendly', level=2)
    document.add_paragraph('Yes' if mobile_friendly else 'No' if mobile_friendly is not None else 'Not Available')
    
    document.save(f"SEO_Report_{url.replace('://', '_').replace('/', '_')}.docx")

def main():
    try:
        url = input("Enter the URL of the webpage: ")
        keywords = input("Enter keywords (comma-separated): ").split(',')
        
        html_content = get_html_content(url)
        
        if html_content:
            page_title, meta_description = get_page_title_and_meta_description(html_content)
            keyword_presence = check_keyword_presence(html_content, keywords)
            page_load_speed = measure_page_load_speed(url)
            mobile_friendly = check_mobile_friendly(url)
            
            generate_report(url, keywords, page_title, meta_description, keyword_presence, page_load_speed, mobile_friendly)
            
            print(f"\nSEO Report generated: SEO_Report_{url.replace('://', '_').replace('/', '_')}.docx")
                
        else:
            print("Unable to fetch HTML. Check the URL and try again.")
    
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

if __name__ == "__main__":
    main()
